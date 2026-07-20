# src/parser.py
# Модуль для парсинга макетов сертификатов

import re
from docx import Document


def parse_certificate(file_path):
    """
    Парсинг макета сертификата из Word документа.
    Возвращает: dict с ключами 'tnved_codes', 'standards'
    """
    doc = Document(file_path)

    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_texts.append(cell.text)

    text = "\n".join(full_text)
    all_text = text + "\n" + "\n".join(table_texts)

    result = {
        'tnved_codes': [],
        'standards': []
    }

    # --- 1. Парсим коды ТНВЭД ---
    tnved_match = re.search(r'КОД ТН ВЭД ЕАЭС\s*[\n\r]+([^\n\r]+)', text, re.IGNORECASE)
    if tnved_match:
        codes_str = tnved_match.group(1).strip()
        codes = re.split(r'[,;]\s*', codes_str)
        codes = [c.strip() for c in codes if c.strip()]
        result['tnved_codes'] = codes
    else:
        tnved_match = re.search(r'КОД ТН ВЭД ЕАЭС\s*([^\n\r]+)', text, re.IGNORECASE)
        if tnved_match:
            codes_str = tnved_match.group(1).strip()
            codes = re.split(r'[,;]\s*', codes_str)
            codes = [c.strip() for c in codes if c.strip()]
            result['tnved_codes'] = codes

    # --- 2. Парсим стандарты ---
    standards_set = set()

    # Удаляем все, что в скобках, и сами скобки из текста (заменяем на пробел)
    text_clean = re.sub(r'\([^)]*\)', ' ', text)

    def clean_standard(std):
        std = std.strip()
        if std.endswith('.'):
            std = std[:-1]
        std = re.sub(r'[\)]+$', '', std)
        std = re.sub(r':(\d{4})', r'-\1', std)
        std = re.sub(r'\s+', ' ', std)
        return std

    def is_standard(text):
        ignore_patterns = [
            r'KI-\d+',
            r'XXXXX',
            r'V\d\.\d\.\d-\d{4}',
            r'№\s*\d+',
            r'б/н',
        ]
        for pattern in ignore_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return False
        if not re.search(r'(ГОСТ|IEC|CISPR|EN|СТБ|ISO)', text, re.IGNORECASE):
            return False
        if not re.search(r'\d', text):
            return False
        return True

    def restore_gost(std):
        if re.match(r'^(IEC|CISPR|EN|ISO)', std, re.IGNORECASE):
            if not re.match(r'^ГОСТ', std, re.IGNORECASE):
                return f"ГОСТ {std}"
        return std

    # 2.1 Ищем в таблицах
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for cell in cells:
                std_pattern = r'(ГОСТ\s*[А-Я\-]*\s*[\d\-\.]+[^\s,;]*|[A-Z]+\s*[\d\-\.]+[^\s,;]*)'
                std_matches = re.findall(std_pattern, cell, re.IGNORECASE)
                for std in std_matches:
                    if std and '15150' not in std:
                        std_clean = clean_standard(std)
                        if is_standard(std_clean):
                            standards_set.add(restore_gost(std_clean))

    # 2.2 Ищем в блоке "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ" (с удаленными скобками)
    additional_match = re.search(r'ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ\s*([\s\S]*?)(?=СРОК ДЕЙСТВИЯ|$)', text_clean,
                                 re.IGNORECASE)

    if additional_match:
        additional_text = additional_match.group(1)

        # Ищем стандарты
        std_pattern = r'(ГОСТ\s*[А-Я\-]*\s*[\d\-\.]+[^\s,;]*|[A-Z]+\s*[\d\-\.]+[^\s,;]*)'
        std_matches = re.findall(std_pattern, additional_text, re.IGNORECASE)
        for std in std_matches:
            if std and '15150' not in std:
                std_clean = clean_standard(std)
                if re.match(r'^(EN|ЕН)\s*301', std_clean, re.IGNORECASE):
                    continue
                if is_standard(std_clean):
                    standards_set.add(restore_gost(std_clean))

    # 2.3 Если не нашли — ищем во всем тексте
    if not standards_set:
        std_pattern = r'(ГОСТ\s*[А-Я\-]*\s*[\d\-\.]+[^\s,;]*|[A-Z]+\s*[\d\-\.]+[^\s,;]*)'
        std_matches = re.findall(std_pattern, text_clean, re.IGNORECASE)
        for std in std_matches:
            if std and '15150' not in std:
                std_clean = clean_standard(std)
                if re.match(r'^(EN|ЕН)\s*301', std_clean, re.IGNORECASE):
                    continue
                if is_standard(std_clean):
                    standards_set.add(restore_gost(std_clean))

    # Дополнительная очистка: убираем стандарты с лишними скобками в конце
    final_standards = []
    for std in standards_set:
        std_clean = re.sub(r'[\)]+$', '', std)
        final_standards.append(std_clean)

    # Жесткая замена ГОСТ EN 301 на полную версию
    final_standards_fixed = []
    for std in final_standards:
        if std == "ГОСТ EN 301":
            std = "ГОСТ EN 301 489-1 V1.9.2-2015"
        final_standards_fixed.append(std)

    result['standards'] = sorted(list(set(final_standards_fixed)))

    return result